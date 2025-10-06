from docxtpl import DocxTemplate
from docx.shared import Inches
import pandas as pd
from datetime import datetime
import matplotlib.pyplot as plt
from docxtpl import InlineImage
from docx import Document
from docx.shared import Pt # <- para tamaño de fuente
from docxtpl import Subdoc
import os
import glob
from io import BytesIO

class InformeHospitalGenerator:
    def __init__(self, template_path):
        """
        Inicializa con plantilla que contiene marcadores {{variable}}
        
        Args:
            template_path: Ruta al archivo .docx con marcadores
            sede: 'Medellin' o 'Rionegro'
        """
        self.doc = DocxTemplate(template_path)
        self.context = {}  # Diccionario con todos los datos
    
    
    def agregar_grafica(self, plt_figure, nombre_marcador):
        """
        Guarda gráfica matplotlib y la referencia para insertar en el marcador
        
        Args:
            plt_figure: Figura de matplotlib
            nombre_marcador: Nombre del marcador en la plantilla, ej: 'grafica_rastreros_1'
        """
        # Guardar imagen temporalmente
        img_path = f'temp_{nombre_marcador}.png'
        plt_figure.savefig(img_path, dpi=150, bbox_inches='tight')
        plt.close(plt_figure)
        
        # Crear objeto InlineImage para docxtpl

        imagen = InlineImage(self.doc, img_path, width=Inches(6))
        
        self.context[nombre_marcador] = imagen
    
    def agregar_contenido_llm(self, contenido, nombre_marcador):
        """
        Agrega contenido generado por LLM a un marcador específico
        
        Args:
            contenido: Texto generado por el LLM
            nombre_marcador: Nombre del marcador, ej: 'analisis_mensual'
        """
        self.context[nombre_marcador] = contenido
    
    def agregar_plot_resultado(self, plot_function, df, nombre_marcador_plot):
        """
        Ejecuta función de plotting y agrega el plot al marcador especificado
        
        Args:
            plot_function: Función que retorna (dataframe, figura), ej: generate_order_area_plot
            df: DataFrame de entrada para la función
            nombre_marcador_plot: Nombre del marcador para el plot, ej: 'preventivos_plot'
        """
        try:
            # Ejecutar función de plotting que retorna (dataframe, figura)
            _, fig = plot_function(df)
            
            # Guardar imagen temporalmente
            img_path = f'temp_{nombre_marcador_plot}.png'
            fig.savefig(img_path, dpi=300, bbox_inches='tight')
            plt.close(fig)
            
            # Crear objeto InlineImage para docxtpl
            
            imagen = InlineImage(self.doc, img_path, width=Inches(6.5))
            
            self.context[nombre_marcador_plot] = imagen
            
        except Exception as e:
            print(f"❌ Error al agregar plot {nombre_marcador_plot}: {e}")
            self.context[nombre_marcador_plot] = f"[Error generando gráfico: {e}]"
    
    def agregar_dataframe_tabla(self, df, nombre_marcador_tabla, table_style='Table Grid'):
        """
        Convierte DataFrame a tabla nativa de Word usando docxtpl
        
        Args:
            df: pandas DataFrame  
            nombre_marcador_tabla: Nombre del marcador para la tabla, ej: 'preventivos_tabla'
            table_style: Estilo de tabla de Word, ej: 'Medium Grid 1 Accent 1'
        """
        try:
            # Para tablas nativas de Word, usar subdoc para crear tabla real
            
            # Crear documento temporal para la tabla
            temp_doc = Document()
            
            # Crear tabla en el documento temporal
            tabla = temp_doc.add_table(rows=1, cols=len(df.columns))
            tabla.style = table_style
            
            # Agregar encabezados
            header_cells = tabla.rows[0].cells
            for i, col_name in enumerate(df.columns):
                header_cells[i].text = str(col_name)
                # Hacer encabezados en negrita
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(6)
            
            # Agregar filas de datos
            for _, row in df.iterrows():
                row_cells = tabla.add_row().cells
                for i, col in enumerate(df.columns):
                    valor = row[col]
                    if pd.isna(valor):
                        row_cells[i].text = '-'
                        for run in row_cells[i].paragraphs[0].runs:
                            run.font.size = Pt(6)
                    elif isinstance(valor, (int, float)):
                        if valor == int(valor):
                            row_cells[i].text = f"{int(valor):,}"
                            for run in row_cells[i].paragraphs[0].runs:
                                run.font.size = Pt(6)
                        else:
                            row_cells[i].text = f"{valor:,.1f}"
                            for run in row_cells[i].paragraphs[0].runs:
                                run.font.size = Pt(6)
                    else:
                        row_cells[i].text = str(valor)
                        for run in row_cells[i].paragraphs[0].runs:
                            run.font.size = Pt(6)
            
            # Guardar tabla temporal
            temp_path = f'temp_tabla_{nombre_marcador_tabla}.docx'
            temp_doc.save(temp_path)
            
            # Crear subdocumento para insertar en el template principal
            
            subdoc = Subdoc(self.doc, temp_path)
            
            self.context[nombre_marcador_tabla] = subdoc
            
            print(f"✅ Tabla Word nativa agregada: {nombre_marcador_tabla} ({len(df)} filas)")
            
        except Exception as e:
            print(f"❌ Error creando tabla nativa, usando formato simple: {e}")
            # Fallback: usar formato de lista simple
            tabla_data = []
            for _, row in df.iterrows():
                fila_dict = {}
                for col in df.columns:
                    valor = row[col]
                    if pd.isna(valor):
                        fila_dict[col] = '-'
                    elif isinstance(valor, (int, float)):
                        if valor == int(valor):
                            fila_dict[col] = f"{int(valor):,}"
                        else:
                            fila_dict[col] = f"{valor:,.1f}"
                    else:
                        fila_dict[col] = str(valor)
                tabla_data.append(fila_dict)
            
            self.context[nombre_marcador_tabla] = {
                'headers': list(df.columns),
                'rows': tabla_data
            }
    
    def agregar_resultado_completo(self, plot_function, df, nombre_marcador_plot, nombre_marcador_tabla):
        """
        Ejecuta función de plotting y agrega tanto el plot como el dataframe a sus respectivos marcadores
        
        Args:
            plot_function: Función que retorna (dataframe, figura), ej: generate_order_area_plot
            df: DataFrame de entrada para la función
            nombre_marcador_plot: Nombre del marcador para el plot, ej: 'preventivos_plot'
            nombre_marcador_tabla: Nombre del marcador para la tabla, ej: 'preventivos_tabla'
        """
        try:
            # Ejecutar función de plotting que retorna (dataframe, figura)
            result_df, fig = plot_function(df)
            
            # Agregar plot
            img_path = f'temp_{nombre_marcador_plot}.png'
            fig.savefig(img_path, dpi=300, bbox_inches='tight')
            plt.close(fig)
            
            
            imagen = InlineImage(self.doc, img_path, width=Inches(6.5))
            self.context[nombre_marcador_plot] = imagen
            
            # Agregar tabla usando el DataFrame resultado
            self.agregar_dataframe_tabla(result_df, nombre_marcador_tabla)
            
            print(f"✅ Agregado plot y tabla: {nombre_marcador_plot}, {nombre_marcador_tabla}")
            
        except Exception as e:
            print(f"❌ Error al agregar resultado completo: {e}")
            self.context[nombre_marcador_plot] = f"[Error generando gráfico: {e}]"
            self.context[nombre_marcador_tabla] = {'headers': [], 'rows': []}
    
    def generar_informe(self, output_path=None, return_buffer=False):
        """
        Renderiza la plantilla con todos los datos y guarda el documento
        
        Args:
            output_path: Ruta donde guardar el archivo (opcional si return_buffer=True)
            return_buffer: Si True, retorna BytesIO buffer en lugar de guardar archivo
            
        Returns:
            BytesIO buffer si return_buffer=True, None en caso contrario
        """
        try:
            # Debug: mostrar las variables del contexto
            print("🔍 Variables en el contexto:")
            for key, value in self.context.items():
                if isinstance(value, dict) and 'headers' in value:
                    print(f"  - {key}: tabla con {len(value['rows'])} filas")
                else:
                    print(f"  - {key}: {type(value).__name__}")
            
            # Renderizar todos los marcadores con los datos del context
            self.doc.render(self.context)
            
            if return_buffer:
                # Guardar en buffer BytesIO para Streamlit
                buffer = BytesIO()
                self.doc.save(buffer)
                buffer.seek(0)
                print(f"✅ Informe generado exitosamente en buffer")
                
                # Limpiar archivos temporales después de guardar
                self._cleanup_temp_files()
                
                return buffer
            else:
                # Guardar documento en archivo
                self.doc.save(output_path)
                print(f"✅ Informe generado exitosamente: {output_path}")
                
                # Limpiar archivos temporales después de guardar
                self._cleanup_temp_files()
                
                return None
            
        except Exception as e:
            print(f"❌ Error al generar informe: {e}")
            print("💡 Revisa la sintaxis de los marcadores en tu plantilla Word")
            print("   Los marcadores deben tener formato: {{nombre_variable}}")
            
            # Limpiar archivos temporales incluso si hay error
            self._cleanup_temp_files()
            
            raise e
    
    def _cleanup_temp_files(self):
        """
        Limpia archivos temporales creados durante la generación del informe
        """
        # Limpiar imágenes temporales
        for temp_file in glob.glob('temp_*.png'):
            try:
                os.remove(temp_file)
            except:
                pass
        # Limpiar tablas temporales
        for temp_file in glob.glob('temp_tabla_*.docx'):
            try:
                os.remove(temp_file)
            except:
                pass
    
    